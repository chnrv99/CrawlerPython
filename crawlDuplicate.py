import time
import requests
from stem import Signal
from stem.control import Controller
from bs4 import BeautifulSoup
import os
import csv
import pandas as pd
import docx
import subprocess
import spacy
import sqlite3

con = sqlite3.connect("Classification.db")
cur = con.cursor()

def addDB():
    url = input("Enter url of classified website")
    classno = int(input("Enter the class number:\n1 - drugs\n2 - weapons\n3- trafficking/fraud"))
    cur.execute("""
    INSERT INTO types(link, classification)
    VALUES (?,?)
    """, (url, classno))
    con.commit ()
    print("Successfully added")
    

keywords_for_drugs = ['alcohol','drug'
    'Ayahuasca',
    'Cannabis', 'Marijuana','Pot','Weed','Central Nervous System Depressants', 'Benzos',
    'Cocaine','Coke','Crack',
    'GHB',
    'Hallucinogens',
    'Heroin',
    'Inhalants',
    'Ketamine',
    'Khat','Kratom',
    'LSD', 'Acid',
    'MDMA','Ecstas','Molly',
    'Mescaline', 'Peyote',
    'Methamphetamine', 'Crystal','Meth',
    'PCP', 'Angel', 'Dust',
    'Prescription' ,'Opioids', 'Oxy','Percs',
    'Stimulants' ,'Speed',
    'Psilocybin' ,'Magic Mushrooms','Shrooms',
    'Rohypnol' ,'Flunitrazepam','Roofies',
    'Salvia',
    'Steroids', 'Anabolic',
    'Synthetic', 'Cannabinoids', 'K2','Spice',
    'Synthetic', 'Cathinones' ,'Bath', 'Salts','Flakka',
    'Tobacco','Nicotine' ,'Vaping']
# keywords_for_fraud = ['credit cards','credit card','stolen accounts','stolen account','passport','fake identity','drivers license','bank','account','accounts','paypal','visa','mastercard',]
keywords_for_guns = ['guns','weapons','gun','weapon','ammo','glock','pistol','sniper','assualt','rifle','EuroGuns','Kal','Eagle','Desert']
keywords_for_traffiking = ['chat','room','traffiking','fourm','credit cards','credit card','stolen accounts','stolen account','passport','fake identity','drivers license','bank','account','accounts','paypal','visa','mastercard','Fourm']

def detect(title):
    print(list(title.split(' ')))
    # if list(title.split(' ')) in keywords_for_drugs:
    #     return 1
    # 1 is for drugs; 2 is for guns; 3 is for human traffiking
    for i in list(title.split( )):
        if i in keywords_for_drugs:
            return 1
        elif i in keywords_for_guns:
            return 2
        elif i in keywords_for_traffiking:
            return 3
        # else:
        #   return 4
    # elif list(title.split(' ')) in keywords_for_guns:
    #     return 2
    # elif list(title.split(' ')) in keywords_for_fraud:
    #     return 3
    # else:
    #     return 4


def drugs(link,title):
    filename = list(title.split(' '))
    filename = filename[0]
    mydoc = docx.Document()
    # if the website has drugs then we print the drugs and details in form of table and write it in word doc
    url=link
    dfs = pd.read_html(url)
    mydoc.add_paragraph('The website belongs to drugs category. The title(of website) and the items sold are: ')
    mydoc.add_paragraph(title)
    for df in dfs:
        # Converting dataframe to list
        li = df.values.tolist()
        # Printing list
        print(li)
        for i in li:
                mydoc.add_paragraph(str(i))      
    mydoc.save('/home/user/Desktop/Crawler/drugs' + filename+'classified_.docx')
    print("Created Word File Successfully")

def weapons(link,title):
    filename = list(title.split(' '))
    filename = filename[0]
    mydoc = docx.Document()
    # if the website has drugs then we print the drugs and details in form of table and write it in word doc
    url=link
    dfs = pd.read_html(url)
    mydoc.add_paragraph('The website belongs to category weapons, the title(of website) and the details of weapons sold are: ')
    mydoc.add_paragraph(title)
    for df in dfs:
        # Converting dataframe to list
        li = df.values.tolist()
        # Printing list
        print(li)
        for i in li:
                mydoc.add_paragraph(str(i))      
    mydoc.save('/home/user/Desktop/Crawler/weapons' + filename+'classified_.docx')

def traffiking(link,title,name):
    age = int(input("Enter age of the child, as an extra information: "))
    blood_group = input("Enter blood group too")
    filename = list(title.split(' '))
    filename = filename[0]
    mydoc = docx.Document()
    # we are checking for any potential human names, and if found we then grab that html element and show it to user
    response = requests.get(link,headers = headers)
    soup = BeautifulSoup(response.text,'html.parser')
    sentence = soup.get_text()

    nlp = spacy.load('en_core_web_sm')
    doc = nlp(sentence)
    
    if sentence.find(name)!=-1: 
        print("Name found! Extra details of the name are:")
        res_index = sentence.find(name,1)
        start_index = res_index-150
        end_index = res_index+150
        print(sentence[start_index:end_index])
        mydoc.add_paragraph(sentence[start_index:end_index])
        mydoc.save('/home/user/Desktop/Crawler/traffiking' + filename+'_Name_classified_.docx')
        # break
    if sentence.find(str(age))!=-1 and sentence.find('old')!=-1:
        print("Age matched! Printing the required data")
        res_index = sentence.find(str(age),1)
        start_index = res_index-150
        end_index = res_index+150
        print(sentence[start_index:end_index])
        mydoc.add_paragraph(sentence[start_index:end_index])
        mydoc.save('/home/user/Desktop/Crawler/traffiking' + filename+'_Age_classified_.docx')
    if sentence.find(str(age))!=-1:
        print("Blood group Matched! Printing the required data")
        res_index = sentence.find(str(age),1)
        start_index = res_index-150
        end_index = res_index+150
        print(sentence[start_index:end_index])
        mydoc.add_paragraph(sentence[start_index:end_index])
        mydoc.save('/home/user/Desktop/Crawler/traffiking' + filename+'B_Group_classified_.docx')
    
    
    
    print("No specific data found, printing all detected names by NLP")
    for ent in doc.ents:
        if ent.label == 'PERSON':
            print(ent.text,ent.label)
    
def enumerate(link):
    # any potential .onion link can be enumerated using nikto
    output = subprocess.run(["nikto", "-h", link])
    print(output)
    choice = input("Do you want the above data in excel format?")
    if choice == "yes":
        filename = link+"_enumeration"
        mydoc = docx.Document()
        mydoc.add_paragraph(output)
        mydoc.save('/home/user/Desktop/Crawler/' + filename+'classified_.docx')


def get_title(link):
    l1=[]
    response = requests.get(link, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    title=''
    print("The title of "+link+" is: ")
    for data in soup.find_all('title'):
        l1.append(data.get_text().strip())
        print(data.get_text().strip())
        data = data.get_text().strip()
    print('\n\n')
    # return l1
    return data








def crawl():
    with Controller.from_port(port=9051) as controller:
        # Set the controller password
        keywords = input("Enter any specific keyword to search: ")
        controller.authenticate(password='CristianoRonaldoCR7')
        num_links_to_crawl = int(input("Enter number of links to crawl: "))
        count =  0
        url = input("Enter the url: ")
        visited = set()
        queue = [url]
    

        while True:
            # Get the next link in the queue
            link = queue.pop(0)


        # Skip the link if it has already been visited
            if link in visited:
                continue

        # Set the new IP address
            controller.signal(Signal.NEWNYM)
            try:
                # Send the request to the URL
                response = requests.get(link, headers=headers)
            # Parse the response
                soup = BeautifulSoup(response.text, 'html.parser')

            # Find all links on the page
                links = soup.find_all('a')

            # Add any links that contain the keywords to the queue
                for a in links:
                    href = a.get('href')
                    if href == None:
                        break
                    # if any(keyword in href for keyword in keywords):
                    if 'http' in href:    
                        queue.append(href)
                    else:
                        queue.append(link + href)
                    if(href==None):
                        continue
                    elif('http' in href or 'https' in href):
                        queue.append(href)
                visited.add(link)
                count+=1
                title = get_title(link)
                classify = detect(title)
                print(classify)
                # classify=3

                if classify == 1:
                    drugs(url,title)
                elif classify == 2:
                    weapons(url,title)
                elif classify == 3:
                    name = input("Its a potential Trafficking website, Enter name to find: ")
                    traffiking(url,title,name)
                if len(visited)==num_links_to_crawl:
                    print("Visited Links: ")
                    for i in range(visited):
                        print(i)
                if len(visited)> num_links_to_crawl:
                    break
                if count>num_links_to_crawl:
                    break
            except(EOFError):
                print("Exception occured at link: ",link)
            

            



user_agent = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.102 Safari/537.36'
headers = {'User-Agent': user_agent}

print("Select your option: \n1 - To crawl a website till a limit is achieved and collect all the links related to that specific keyword\n2 - To classify the given website\n3 - To check for child details in a given website\n4 - To enumerate a given .onion link \n5 - To add a link to database\n\nSelect your option:") 

option = int(input())

# To crawl a website till a limit is achieved and collect all the links related to that specific keyword
if option == 1:
    crawl()
# To classify the given website
elif option == 2:
    url = input("Enter URL to be detected: ")
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    for data in soup.find_all('title'):
        data = data.get_text().strip()
    classify = detect(data)
    print("1 - Drugs\n2 - Weapons\n3 - Human Traffiking/Fraud: ")
    print(classify)
    if classify == 1 or classify == 2:
        drugs(url,data)
    elif classify == 3:
        name = input("Enter child name: ")
        traffiking(url,data,name)


# To check for child details
elif option == 3:
    url = input("Enter link: ")
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    # title=''
    for data in soup.find_all('title'):
        data = data.get_text().strip()
    # weapons(url,data)
    name = input("Enter the baby name to be searched: ")
    traffiking(url,data,name)
    
    

# to enumerate a given link
elif option == 4:
    url = input("Enter URL: ")
    enumerate(url)

# adding a link and its classification to sqlite db
elif option == 5:
    addDB()







# Drug link : http://6hzbfxpnsdo4bkplp5uojidkibswevsz3cfpdynih3qvfr24t5qlkcyd.onion/
# Weapon link : 
# baby link : https://nameberry.com/
# 4th your choice
# adding to database